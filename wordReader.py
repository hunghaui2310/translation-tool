import docx

def translate_text_in_docx(file_path, source, target):
    doc = docx.Document("input.docx")
    for p in doc.paragraphs:
        # Replace the old text with the new text
        # p.text = p.text.replace(old_text, new_text)
        print(p.text)

    # Save the updated document
    doc.save('output.docx')

if __name__ == '__main__':
    doc = docx.Document("input.docx")

    # Loop through each paragraph in the document
    for p in doc.paragraphs:
        # Replace the old text with the new text
        # p.text = p.text.replace(old_text, new_text)
        print(p.text)

    # Save the updated document
    doc.save('output.docx')